"""Serviço de Geração Automática de Laudos Periciais.

Fluxo:
  1. build_placeholder_context() — coleta dados do DB (device, target, op, hashes, fotos)
  2. generate_docx()             — carrega template DOCX, substitui placeholders, insere imagens
  3. convert_to_pdf()            — converte DOCX → PDF via LibreOffice headless (subprocess)
"""
import io
import logging
import os
import re
import subprocess
import tempfile
import uuid
from datetime import date
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Inches
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.custody_model import CustodyMovement
from app.models.device_model import Device
from app.models.hash_model import IntegrityHash
from app.models.operation_model import Operation
from app.models.photo_model import DevicePhoto
from app.models.target_model import Target

logger = logging.getLogger(__name__)

# ── Mapa de placeholders ─────────────────────────────────────────────
PLACEHOLDER_MAP = {
    # ── Documento / Laudo ─────────────────────────────────────────
    "{{NUMERO_LAUDO}}": "report_number",
    "{{NUMERO_DOCUMENTO}}": "report_number",
    "{{DATA_EMISSAO}}": "emission_date",
    "{{PERITO}}": "expert_name",
    "{{RESPONSAVEL}}": "expert_name",
    "{{OBSERVACOES}}": "observations",
    # ── Operação ─────────────────────────────────────────────────
    "{{OPERACAO}}": "operation_name",
    "{{NOME_OPERACAO}}": "operation_name",
    "{{NUMERO_PROCEDIMENTO}}": "procedure_number",
    "{{PROCEDIMENTO}}": "procedure_number",
    "{{UNIDADE_RESPONSAVEL}}": "responsible_unit",
    "{{DATA_INICIO_OPERACAO}}": "start_date",
    "{{DATA_ENCERRAMENTO_OPERACAO}}": "end_date",
    "{{STATUS_OPERACAO}}": "operation_status",
    "{{TOTAL_ALVOS}}": "total_targets",
    "{{TOTAL_DISPOSITIVOS}}": "total_devices",
    # ── Alvo ──────────────────────────────────────────────────────
    "{{ALVO}}": "target_name",
    "{{NOME_ALVO}}": "target_name",
    "{{CPF_ALVO}}": "target_cpf",
    "{{CPF}}": "target_cpf",
    "{{RG_ALVO}}": "target_rg",
    "{{RG}}": "target_rg",
    "{{APELIDO_ALVO}}": "target_nickname",
    "{{APELIDO}}": "target_nickname",
    "{{DATA_NASCIMENTO}}": "target_birth_date",
    "{{ENDERECO_ALVO}}": "target_address",
    "{{ENDERECO}}": "target_address",
    "{{NOME_SOCIAL}}": "target_social_name",
    "{{TIPO_PESSOA}}": "target_person_type",
    "{{OBSERVACOES_ALVO}}": "target_observations",
    # ── Dispositivo ───────────────────────────────────────────────
    "{{DISPOSITIVO}}": "device_type",
    "{{TIPO_DISPOSITIVO}}": "device_type",
    "{{MARCA}}": "brand",
    "{{MODELO}}": "model",
    "{{NUMERO_EVIDENCIA}}": "evidence_number",
    "{{EVIDENCIA}}": "evidence_number",
    "{{NUMERO_LACRE}}": "seal_number",
    "{{LACRE}}": "seal_number",
    "{{COR}}": "color",
    "{{DATA_APREENSAO}}": "seizure_date",
    "{{LOCAL_APREENSAO}}": "seizure_location",
    "{{OBSERVACOES_APREENSAO}}": "seizure_observations",
    "{{STATUS_DISPOSITIVO}}": "device_status",
    # ── Dados técnicos do dispositivo ────────────────────────────────
    "{{IMEI}}": "imei",
    "{{IMEI1}}": "imei",
    "{{IMEI2}}": "imei2",
    "{{ICCID}}": "iccid",
    "{{TELEFONE}}": "phone_number",
    "{{SERIAL}}": "serial_number",
    "{{NUMERO_SERIE}}": "serial_number",
    "{{SISTEMA_OPERACIONAL}}": "os",
    "{{SO}}": "os",
    "{{CAPACIDADE}}": "storage_capacity",
    "{{ARMAZENAMENTO}}": "storage_capacity",
    "{{MEMORIA_RAM}}": "ram",
    "{{RAM}}": "ram",
    "{{PROCESSADOR}}": "processor",
    "{{CPU}}": "processor",
    "{{INTERFACE}}": "interface",
    # ── Hashes de integridade ──────────────────────────────────────
    "{{HASH_MD5}}": "hash_md5",
    "{{HASH_SHA1}}": "hash_sha1",
    "{{HASH_SHA256}}": "hash_sha256",
    "{{ARQUIVO_EXTRACAO}}": "source_file",
    "{{ARQUIVO_ORIGEM}}": "source_file",
    # ── Custódia ────────────────────────────────────────────────────
    "{{DATA_ANALISE}}": "analysis_start_date",
    "{{DATA_INICIO_ANALISE}}": "analysis_start_date",
    "{{DATA_FIM_ANALISE}}": "analysis_end_date",
    "{{SETOR_ORIGEM}}": "custody_origin",
    "{{SETOR_DESTINO}}": "custody_destination",
    "{{RESPONSAVEL_CUSTODIA}}": "custody_responsible",
    "{{TIPO_MOVIMENTACAO}}": "custody_movement_type",
    "{{RAZAO_MOVIMENTACAO}}": "custody_reason",
    "{{TOTAL_MOVIMENTACOES}}": "total_movements",
    # ── Contagens ───────────────────────────────────────────────────
    "{{TOTAL_FOTOS}}": "photos_count",
    "{{QTDE_FOTOS}}": "photos_count",
}

# Placeholders de imagem — mapeados por categoria de foto
IMAGE_PLACEHOLDER_MAP = {
    "{{FOTO_DISPOSITIVO}}": "front",
    "{{FOTO_FRONTAL}}": "front",
    "{{FOTO_TRASEIRA}}": "back",
    "{{FOTO_LACRE}}": "seal",
    "{{FOTO_SERIAL}}": "serial_number",
    "{{FOTO_IMEI}}": "imei",
    "{{FOTO_APREENSAO}}": "evidence_state",
}


# ── Coleta de dados (context) ─────────────────────────────────────

async def build_placeholder_context(
    session: AsyncSession,
    device_id: uuid.UUID,
    report_number: str,
    expert_name: str | None = None,
    emission_date: date | None = None,
    observations: str | None = None,
) -> dict[str, Any]:
    """Busca todos os dados do DB e monta o contexto de substituição de placeholders."""

    # Device
    device_result = await session.execute(
        select(Device).where(Device.id == device_id, Device.deleted_at.is_(None))
    )
    device: Device | None = device_result.scalar_one_or_none()
    if not device:
        raise ValueError(f"Dispositivo {device_id} não encontrado.")

    extra = device.extra_data or {}

    # Target
    target: Target | None = None
    if device.target_id:
        target_result = await session.execute(
            select(Target).where(Target.id == device.target_id, Target.deleted_at.is_(None))
        )
        target = target_result.scalar_one_or_none()

    # Operation
    operation: Operation | None = None
    if device.operation_id:
        op_result = await session.execute(
            select(Operation).where(Operation.id == device.operation_id, Operation.deleted_at.is_(None))
        )
        operation = op_result.scalar_one_or_none()

    # Hashes — pega o mais recente
    hash_result = await session.execute(
        select(IntegrityHash)
        .where(IntegrityHash.device_id == device_id)
        .order_by(IntegrityHash.calculated_at.desc())
    )
    latest_hash: IntegrityHash | None = hash_result.scalars().first()

    # Custódia — data de início de análise
    custody_result = await session.execute(
        select(CustodyMovement)
        .where(
            CustodyMovement.device_id == device_id,
            CustodyMovement.movement_type == "analysis_start",
        )
        .order_by(CustodyMovement.movement_date.asc())
    )
    analysis_movement: CustodyMovement | None = custody_result.scalars().first()

    # Fotos
    photos_result = await session.execute(
        select(DevicePhoto).where(
            DevicePhoto.device_id == device_id,
            DevicePhoto.deleted_at.is_(None),
        )
    )
    photos: list[DevicePhoto] = list(photos_result.scalars().all())

    def fmt_date(d: Any) -> str | None:
        if d is None:
            return None
        if isinstance(d, date):
            return d.strftime("%d/%m/%Y")
        return str(d)

    context: dict[str, Any] = {
        # Laudo / Documento
        "report_number": report_number,
        "expert_name": expert_name or "",
        "emission_date": fmt_date(emission_date) or "",
        "observations": observations or "",
        # Device
        "evidence_number": device.evidence_number or "",
        "seal_number": device.seal_number or "",
        "device_type": device.device_type or "",
        "device_status": device.status or "",
        "brand": device.brand or "",
        "model": device.model or "",
        "serial_number": device.serial_number or "",
        "color": device.color or "",
        "seizure_date": fmt_date(device.seizure_date) or "",
        "seizure_location": device.seizure_location or "",
        "seizure_observations": device.seizure_observations or "",
        # Device extra_data (smartphone, tablet, etc.)
        "imei": extra.get("imei") or extra.get("imei1") or "",
        "imei2": extra.get("imei2") or "",
        "iccid": extra.get("iccid") or "",
        "phone_number": extra.get("phone_number") or "",
        "os": extra.get("os") or extra.get("operating_system") or "",
        "storage_capacity": extra.get("storage_capacity") or extra.get("capacity") or "",
        "ram": extra.get("ram") or "",
        "processor": extra.get("processor") or "",
        "interface": extra.get("interface") or "",
        # Target
        "target_name": target.full_name if target else "",
        "target_cpf": target.cpf if target else "",
        "target_rg": target.rg if target else "",
        "target_nickname": target.nickname if target else "",
        "target_birth_date": fmt_date(target.birth_date) if target else "",
        "target_address": target.address if target else "",
        "target_social_name": target.social_name if target else "",
        "target_person_type": ("Pessoa Física" if target and target.person_type == "individual" else "Pessoa Jurídica") if target else "",
        "target_observations": target.observations if target else "",
        # Operation
        "operation_name": operation.name if operation else "",
        "procedure_number": operation.procedure_number if operation else "",
        "responsible_unit": operation.responsible_unit if operation else "",
        "start_date": fmt_date(operation.start_date) if operation else "",
        "end_date": fmt_date(operation.end_date) if operation else "",
        "operation_status": operation.status if operation else "",
        "total_targets": "",
        "total_devices": "",
        # Hashes
        "hash_md5": latest_hash.md5 if latest_hash else "",
        "hash_sha1": latest_hash.sha1 if latest_hash else "",
        "hash_sha256": latest_hash.sha256 if latest_hash else "",
        "source_file": latest_hash.source_file if latest_hash else "",
        # Custódia
        "analysis_start_date": (
            fmt_date(analysis_movement.movement_date) if analysis_movement else ""
        ),
        "analysis_end_date": "",
        "custody_origin": "",
        "custody_destination": "",
        "custody_responsible": "",
        "custody_movement_type": "",
        "custody_reason": "",
        "total_movements": "",
        # Photos
        "photos_count": len(photos),
        "_photos": photos,  # lista interna para inserção de imagens
        "_device": device,
    }

    # Preenche dados de custódia a partir de todos os movimentos
    all_custody_result = await session.execute(
        select(CustodyMovement)
        .where(CustodyMovement.device_id == device_id)
        .order_by(CustodyMovement.movement_date.asc())
    )
    all_movements: list[CustodyMovement] = list(all_custody_result.scalars().all())
    context["total_movements"] = str(len(all_movements))

    analysis_end = next((m for m in reversed(all_movements) if m.movement_type == "analysis_end"), None)
    if analysis_end:
        context["analysis_end_date"] = fmt_date(analysis_end.movement_date) or ""

    last_movement = all_movements[-1] if all_movements else None
    if last_movement:
        context["custody_origin"] = last_movement.origin_sector or ""
        context["custody_destination"] = last_movement.destination_sector or ""
        context["custody_responsible"] = last_movement.responsible_name or ""
        context["custody_movement_type"] = last_movement.movement_type or ""
        context["custody_reason"] = last_movement.reason or ""

    return context


# ── Contexto para Operação ────────────────────────────────────────

async def build_operation_context(
    session: AsyncSession,
    operation_id: uuid.UUID,
    report_number: str,
    expert_name: str | None = None,
    emission_date: date | None = None,
    observations: str | None = None,
) -> dict[str, Any]:
    """Monta o contexto de substituição de placeholders a partir de uma Operação."""
    from app.models.target_model import Target
    from app.models.device_model import Device
    from sqlalchemy import func

    # Operation
    op_result = await session.execute(
        select(Operation).where(Operation.id == operation_id, Operation.deleted_at.is_(None))
    )
    operation: Operation | None = op_result.scalar_one_or_none()
    if not operation:
        raise ValueError(f"Operação {operation_id} não encontrada.")

    # Contagem de alvos e dispositivos
    targets_count = (await session.execute(
        select(func.count()).select_from(Target).where(
            Target.operation_id == operation_id,
            Target.deleted_at.is_(None),
        )
    )).scalar_one()

    devices_count = (await session.execute(
        select(func.count()).select_from(Device).where(
            Device.operation_id == operation_id,
            Device.deleted_at.is_(None),
        )
    )).scalar_one()

    def fmt_date(d: Any) -> str | None:
        if d is None:
            return None
        if isinstance(d, date):
            return d.strftime("%d/%m/%Y")
        return str(d)

    context: dict[str, Any] = {
        # Documento
        "report_number": report_number,
        "expert_name": expert_name or "",
        "emission_date": fmt_date(emission_date) or "",
        "observations": observations or "",
        # Operação
        "operation_name": operation.name or "",
        "procedure_number": operation.procedure_number or "",
        "responsible_unit": operation.responsible_unit or "",
        "start_date": fmt_date(operation.start_date) or "",
        "end_date": fmt_date(operation.end_date) or "",
        "operation_status": operation.status or "",
        "total_targets": str(targets_count),
        "total_devices": str(devices_count),
        # Campos de dispositivo — deixados vazios para templates genéricos
        "evidence_number": "",
        "seal_number": "",
        "device_type": "",
        "brand": "",
        "model": "",
        "serial_number": "",
        "color": "",
        "seizure_date": "",
        "seizure_location": "",
        "seizure_observations": "",
        "imei": "",
        "os": "",
        "storage_capacity": "",
        "ram": "",
        "processor": "",
        "target_name": "",
        "target_cpf": "",
        "hash_md5": "",
        "hash_sha1": "",
        "hash_sha256": "",
        "source_file": "",
        "analysis_start_date": "",
        "photos_count": 0,
        "_photos": [],
    }
    return context


# ── Contexto para Alvo ────────────────────────────────────────────

async def build_target_context(
    session: AsyncSession,
    target_id: uuid.UUID,
    report_number: str,
    expert_name: str | None = None,
    emission_date: date | None = None,
    observations: str | None = None,
) -> dict[str, Any]:
    """Monta o contexto de substituição de placeholders a partir de um Alvo."""
    from app.models.device_model import Device
    from sqlalchemy import func

    # Target
    target_result = await session.execute(
        select(Target).where(Target.id == target_id, Target.deleted_at.is_(None))
    )
    target: Target | None = target_result.scalar_one_or_none()
    if not target:
        raise ValueError(f"Alvo {target_id} não encontrado.")

    # Operation
    operation: Operation | None = None
    if target.operation_id:
        op_result = await session.execute(
            select(Operation).where(
                Operation.id == target.operation_id, Operation.deleted_at.is_(None)
            )
        )
        operation = op_result.scalar_one_or_none()

    # Dispositivos do alvo
    devices_count = (await session.execute(
        select(func.count()).select_from(Device).where(
            Device.target_id == target_id,
            Device.deleted_at.is_(None),
        )
    )).scalar_one()

    def fmt_date(d: Any) -> str | None:
        if d is None:
            return None
        if isinstance(d, date):
            return d.strftime("%d/%m/%Y")
        return str(d)

    context: dict[str, Any] = {
        # Documento
        "report_number": report_number,
        "expert_name": expert_name or "",
        "emission_date": fmt_date(emission_date) or "",
        "observations": observations or "",
        # Alvo
        "target_name": target.full_name or "",
        "target_cpf": target.cpf or "",
        "target_rg": target.rg or "",
        "target_nickname": target.nickname or "",
        "target_birth_date": fmt_date(target.birth_date) if hasattr(target, 'birth_date') else "",
        "target_address": target.address if hasattr(target, 'address') else "",
        "total_devices": str(devices_count),
        # Operação
        "operation_name": operation.name if operation else "",
        "procedure_number": operation.procedure_number if operation else "",
        "responsible_unit": operation.responsible_unit if operation else "",
        "start_date": fmt_date(operation.start_date) if operation else "",
        # Campos de dispositivo — deixados vazios
        "evidence_number": "",
        "seal_number": "",
        "device_type": "",
        "brand": "",
        "model": "",
        "serial_number": "",
        "color": "",
        "seizure_date": "",
        "seizure_location": "",
        "seizure_observations": "",
        "imei": "",
        "os": "",
        "storage_capacity": "",
        "ram": "",
        "processor": "",
        "hash_md5": "",
        "hash_sha1": "",
        "hash_sha256": "",
        "source_file": "",
        "analysis_start_date": "",
        "photos_count": 0,
        "_photos": [],
    }
    return context


# ── Geração DOCX ──────────────────────────────────────────────────

def _replace_text_in_paragraph(para: Any, context: dict[str, Any]) -> None:
    """Substitui placeholders dentro de um parágrafo preservando a formatação dos runs."""
    # Primeiro constrói o texto completo do parágrafo
    full_text = "".join(run.text for run in para.runs)
    has_placeholder = any(ph in full_text for ph in PLACEHOLDER_MAP.keys())
    if not has_placeholder:
        return

    # Para cada run, substitui os placeholders de texto
    for run in para.runs:
        for placeholder, key in PLACEHOLDER_MAP.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(context.get(key) or ""))


def _replace_text_in_table(table: Any, context: dict[str, Any]) -> None:
    """Substitui placeholders em todas as células da tabela."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_text_in_paragraph(para, context)


def _get_photo_bytes_by_category(
    photos: list[DevicePhoto],
    category: str,
    photo_bytes_map: dict[str, bytes],
) -> bytes | None:
    """Retorna os bytes da primeira foto com a categoria especificada."""
    for photo in photos:
        if photo.category == category and photo.file_path in photo_bytes_map:
            return photo_bytes_map[photo.file_path]
    return None


def _insert_image_in_paragraph(para: Any, img_bytes: bytes, width_inches: float = 3.0) -> None:
    """Substitui o conteúdo do parágrafo por uma imagem inline."""
    from docx.oxml.ns import qn
    # Limpa todos os runs do parágrafo
    for run in para.runs:
        run.text = ""
    # Adiciona a imagem no primeiro run (ou cria um novo)
    if not para.runs:
        run = para.add_run()
    else:
        run = para.runs[0]
    img_stream = io.BytesIO(img_bytes)
    run.add_picture(img_stream, width=Inches(width_inches))


def generate_docx(
    template_bytes: bytes,
    context: dict[str, Any],
    photo_bytes_map: dict[str, bytes] | None = None,
) -> bytes:
    """
    Carrega o template DOCX, substitui todos os placeholders de texto e insere imagens.
    Retorna o DOCX gerado como bytes.
    """
    photo_bytes_map = photo_bytes_map or {}
    photos: list[DevicePhoto] = context.get("_photos", [])

    doc = DocxDocument(io.BytesIO(template_bytes))

    # 1. Substituir texto em parágrafos do corpo
    for para in doc.paragraphs:
        # Verifica se é um placeholder de imagem
        full_text = "".join(run.text for run in para.runs).strip()
        image_matched = False
        for img_ph, category in IMAGE_PLACEHOLDER_MAP.items():
            if img_ph in full_text:
                img_bytes = _get_photo_bytes_by_category(photos, category, photo_bytes_map)
                if img_bytes:
                    _insert_image_in_paragraph(para, img_bytes)
                else:
                    # Remove o placeholder sem imagem
                    for run in para.runs:
                        run.text = run.text.replace(img_ph, f"[Imagem não disponível: {img_ph}]")
                image_matched = True
                break

        if not image_matched:
            _replace_text_in_paragraph(para, context)

    # 2. Substituir texto em tabelas
    for table in doc.tables:
        _replace_text_in_table(table, context)

    # 3. Substituir em cabeçalhos e rodapés
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                _replace_text_in_paragraph(para, context)
        if section.footer:
            for para in section.footer.paragraphs:
                _replace_text_in_paragraph(para, context)

    # 4. Salva em bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ── Conversão DOCX → PDF via LibreOffice headless ─────────────────

def convert_to_pdf(docx_bytes: bytes) -> bytes | None:
    """
    Converte um DOCX (bytes) para PDF usando LibreOffice headless.
    Retorna os bytes do PDF ou None em caso de falha.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "laudo.docx")
        pdf_path = os.path.join(tmpdir, "laudo.pdf")

        # Escreve o DOCX em disco
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # Chama LibreOffice headless
        try:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    docx_path,
                ],
                capture_output=True,
                timeout=120,  # 2 minutos de timeout
                text=True,
            )
            if result.returncode != 0:
                logger.error(
                    "LibreOffice conversion failed. stdout=%s, stderr=%s",
                    result.stdout,
                    result.stderr,
                )
                return None
        except subprocess.TimeoutExpired:
            logger.error("LibreOffice conversion timed out after 120s.")
            return None
        except FileNotFoundError:
            logger.error("LibreOffice não encontrado. Certifique-se que está instalado no container.")
            return None

        # Lê o PDF gerado
        if not os.path.exists(pdf_path):
            logger.error("PDF não foi gerado: %s", pdf_path)
            return None

        with open(pdf_path, "rb") as f:
            return f.read()
